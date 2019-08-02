from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import pandas as pd
import numpy as np
import pymysql
import json
import datetime as dt

from pandas.plotting import register_matplotlib_converters
register_matplotlib_converters()

#matplotlib.use('tkagg')
#https://codeyarns.com/2016/03/14/matplotlib-plot-is-not-displayed-in-window/

conn = pymysql.connect(host='host',
                             user='user',
                             password='pass',
                             db='db',
                             charset='utf8') #,
#cursorclass=pymysql.cursors.DictCursor)

with open("devs_A_state.json") as json_file:
    try:
        json_data = json.load(json_file)
    except:
        print("Error loading JSON file.")

dev_state_by_sids = json_data['dev_state']
#dev_state_by_sids = dev_state_s3["dev_state"]
stations = dev_state_by_sids.keys()


document = Document()

#document.add_heading('Document Title', 0)
#p = document.add_paragraph('A plain paragraph having some ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

heading1 = document.add_picture('bg_logo.png', width=Inches(2))
heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_heading('BluGraph Technologies Pte Ltd', level=1)
document.add_heading('Supply of Water Level Data at Waterways (Fifth Contract)', level=1)
document.add_heading('Monthly Report', level=1)
document.add_heading('May  2019', level=1)
document.add_heading('System - A', level=1)
#document.add_paragraph('Intense quote', style='Intense Quote')
document.add_page_break()

for sid in stations:
    print("Processing.. ", sid)
    # Get station details.
    # Query DB
    df2 = pd.read_sql('select station_name, copelevel, invertlevel, height, operationlevel, criticallevel from bwl_station where station_id=\"' + sid + '\"', conn)
    print(df2.head())
    station_name = df2.iloc[0]['station_name']
    invertlevel = float(df2.iloc[0]['invertlevel'])
    copelevel = float(df2.iloc[0]['copelevel'])
    operationlevel = float(df2.iloc[0]['operationlevel'])
    print(station_name)
    #cursor = conn.cursor()
    #cursor.execute('select Name, Continent, Population, LifeExpectancy, GNP from country')
    #rows = cursor.fetchall()

    #df = pd.DataFrame(rows)

    #print(str(rows)[0:300])

    document.add_heading('Station ID: ' + sid, level=2)
    document.add_heading('Station Name: ' + station_name, level=2)

    # Get station details.
    pc50_val = invertlevel + (copelevel - invertlevel) / 2
    pc75_val = invertlevel + (copelevel - invertlevel) * 75/100
    pc90_val = invertlevel + (copelevel - invertlevel) * 90/100
    pc100_val = copelevel
    cl_val = float(df2.iloc[0]['criticallevel'])
    # 60?

#    p = document.add_paragraph('A plain paragraph having some ')
#    p.add_run('bold').bold = True
#    p.add_run(' and some ')
#    p.add_run('italic.').italic = True

    # Generate and save monthly chart.
    # Query DB
    st = '2019-06-01 00:00:00'
    et = '2019-06-30 23:59:00'
    df1 = pd.read_sql('select datetime, waterlever_mrl, maintenance_status from raw_data where station_id=\"' + sid + '\" and datetime between \"' + st + '\" and \"' + et + '\" order by datetime desc', 
                   conn,
                   parse_dates = ['datetime'], 
                   index_col = ['datetime'])
    #df1 = pd.read_csv('raw_data_CWS001_mar2019.csv', parse_dates = ['dt'], index_col = ['dt'])

    df1['waterlever_mrl'] = df1['waterlever_mrl'].astype(float)
    #df1['maintenance_status'] = df1['waterlever_mrl'].astype(int)
    print(df1.dtypes)
    print(df1.head())

    fig, ax = plt.subplots(figsize=(8, 3))
    ax.plot(df1.index.values, df1['waterlever_mrl'])
    plt.setp(ax.get_xticklabels(), rotation=0)
    #50% line
    ax.axhline(pc50_val, color='blue')
    x_text_annotation = dt.datetime(2019, 6, 2)
    plt.text(x=x_text_annotation, y=pc50_val+0.05,s='50%',rotation=0)
    # 75% line
    ax.axhline(pc75_val, color='green')
    x_text_annotation = dt.datetime(2019, 6, 5)
    plt.text(x=x_text_annotation, y=pc75_val+0.05,s='75%',rotation=0)
    # 90% line
    ax.axhline(pc90_val, color='yellow')
    x_text_annotation = dt.datetime(2019, 6, 7)
    plt.text(x=x_text_annotation, y=pc90_val+0.05,s='90%',rotation=0)
    # 100% line
    ax.axhline(pc100_val, color='purple')
    x_text_annotation = dt.datetime(2019, 6, 9)
    plt.text(x=x_text_annotation, y=pc100_val+0.05,s='100%',rotation=0)
    # critical value line
    ax.axhline(cl_val, color='red')
    x_text_annotation = dt.datetime(2019, 6, 11)
    plt.text(x=x_text_annotation, y=cl_val+0.05,s='critical level',rotation=0)

    #ax.set_xticks(df1.index)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
    #ax.xaxis.set_minor_formatter(mdates.DateFormatter("%Y-%m"))
    #plt.xticks(rotation=90)  

    ax.pcolorfast(ax.get_xlim(), ax.get_ylim(),
              df1['maintenance_status'].values[np.newaxis],
              cmap='gray_r', alpha=0.3)
              #cmap='RdYlGn', alpha=0.3)

    plt.savefig('monthly-chart.png', bbox_inches='tight')

    document.add_picture('monthly-chart.png', width=Inches(6.5))

    # Calculate Max and Min from query results. 
    # Remove data in maintenance mode.
    is_active = df1['maintenance_status'] == 0
    df1_active = df1[is_active]
    max_val = df1_active['waterlever_mrl'].max()
    min_val = df1_active['waterlever_mrl'].min()

    paragraph1 = document.add_paragraph(
        #'Cope Level(mRI): ' + str(copelevel), style="BodyText", breakbefore=False, jc='left'
        'Cope Level(mRL): ' + str(copelevel), style="BodyText"
        #'Cope Level(mRI): ' + cope_level, style='List Bullet'
    )
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(
        'Operational Level(mRL): ' + str(operationlevel), style='BodyText'
    )
    document.add_paragraph(
        'Invert Level(mRL): ' + str(invertlevel), style='BodyText'
    )
    document.add_paragraph(
        'Max: ' + str(max_val), style='BodyText'
    )
    document.add_paragraph(
        'Min: ' + str(min_val), style='BodyText'
    )

    document.add_page_break()

document.save('System-A.docx')
